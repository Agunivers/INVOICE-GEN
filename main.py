import os
import customtkinter as ctk
import random
import datetime
from tkinter import messagebox, filedialog
from docx2pdf import convert
import docx
from openpyxl import Workbook, load_workbook

class Application(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("INVOICE GENERATOR")
        self.geometry("600x400")
        self.resizable(False, False)

        # Set theme
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")

        # Create UI elements
        self.create_widgets()

        # Invoice Number Management
        self.sequence_id = "INV"
        self.current_invoice_number = None

        self.mainloop()

    def create_widgets(self):
        """Creates and places widgets."""
        self.name_entry = ctk.CTkEntry(self, placeholder_text="Enter your Name", width=250, height=35)
        self.number_entry = ctk.CTkEntry(self, placeholder_text="Enter your Number", width=250, height=35)
        self.email_entry = ctk.CTkEntry(self, placeholder_text="Enter your Email", width=250, height=35)

        self.create_button = ctk.CTkButton(self, text="Create Invoice", command=self.invoice, width=130, height=40)
        self.save_button = ctk.CTkButton(self, text="Save to Excel", command=self.save_to_excel, width=130, height=40)

        self.name_entry.place(x=180, y=50)
        self.number_entry.place(x=180, y=100)
        self.email_entry.place(x=180, y=150)
        self.create_button.place(x=150, y=220)
        self.save_button.place(x=320, y=220)

    def generate_invoice_number(self):
        """Generates a unique invoice number."""
        if self.current_invoice_number is None:
            self.current_invoice_number = random.randint(1000, 9999)
        return self.current_invoice_number

    def invoice(self):

        doc = docx.Document("invoice.docx")
        invoice_number = f"{self.sequence_id}-{datetime.datetime.now().strftime('%Y')}-{self.generate_invoice_number()}"

        replacements = {
            "[name]": self.name_entry.get().upper(),
            "[number]": self.number_entry.get(),
            "[email]": self.email_entry.get(),
            "[date]": datetime.datetime.now().strftime("%d/%m/%Y"),
            "[time]": datetime.datetime.now().strftime("%H:%M:%S"),
            "[inv]": self.sequence_id,
            "[year]": datetime.datetime.now().strftime("%Y"),
            "[num]": str(self.current_invoice_number)
        }

        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                paragraph.text = paragraph.text.replace(old_text, new_text)

        temp_doc_path = "temp_invoice.docx"
        doc.save(temp_doc_path)

        save_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF document', '*.pdf')])
        if not save_path:
            messagebox.showinfo("Cancelled", "Invoice creation cancelled.")
            return

        try:
            convert(temp_doc_path, os.path.dirname(save_path))
            os.rename(os.path.join(os.path.dirname(save_path), "temp_invoice.pdf"), save_path)
            os.remove(temp_doc_path)
            messagebox.showinfo("Success", f"Invoice created successfully at:\n{save_path}")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")
            os.remove(temp_doc_path)

        self.save_to_excel()

    def save_to_excel(self):
        """Saves invoice data to an Excel file."""
        name = self.name_entry.get().strip()
        number = self.number_entry.get().strip()
        email = self.email_entry.get().strip()
        invoice_number = f"{self.sequence_id}-{datetime.datetime.now().strftime('%Y')}-{self.current_invoice_number}"

        if not name or not number or not email:
            messagebox.showerror("Validation Error", "All fields are required!")
            return

        file_path = "USER_DATA.xlsx"
        try:
            if os.path.exists(file_path):
                workbook = load_workbook(file_path)
                sheet = workbook.active
            else:
                workbook = Workbook()
                sheet = workbook.active
                sheet.append(["Name", "Number", "Email", "Date", "Time", "Invoice Number"])

            sheet.append([name, number, email,
                          datetime.datetime.now().strftime("%d/%m/%Y"),
                          datetime.datetime.now().strftime("%H:%M:%S"),
                          invoice_number])

            workbook.save(file_path)
            workbook.close()
            messagebox.showinfo("Success", f"Data saved to {file_path}!")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while saving to Excel: {e}")

if __name__ == "__main__":
    app = Application()
