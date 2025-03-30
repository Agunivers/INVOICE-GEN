import os
import random
import datetime
import customtkinter as ctk
from tkinter import messagebox, filedialog
from docx2pdf import convert
import docx
from openpyxl import Workbook, load_workbook


class Application():
    def __init__(self):
        ctk.set_appearance_mode("dark")  # Set theme
        ctk.set_default_color_theme("blue")

        self.root = ctk.CTk()
        self.root.title("Excel to PDF Converter")
        self.root.geometry("500x300")
        self.root.resizable(False, False)

        # Create UI Elements
        self.name_entry = ctk.CTkEntry(self.root, placeholder_text="Enter your name", width=250, height=35)
        self.number_entry = ctk.CTkEntry(self.root, placeholder_text="Enter your number", width=250, height=35)
        self.email_entry = ctk.CTkEntry(self.root, placeholder_text="Enter your email", width=250, height=35)

        self.create_button = ctk.CTkButton(self.root, text="Create Invoice", command=self.invoice, width=130, height=40)
        self.save_button = ctk.CTkButton(self.root, text="Save to Excel", command=self.save_to_excel, width=130,
                                         height=40)

        # Place UI Elements
        self.name_entry.place(x=125, y=50)
        self.number_entry.place(x=125, y=100)
        self.email_entry.place(x=125, y=150)
        self.create_button.place(x=100, y=220)
        self.save_button.place(x=270, y=220)

        # Initialize Variables
        self.sequence_id = "INV"
        self.sequence_number = set()
        self.current_invoice_number = None

        self.root.mainloop()

    def number(self):
        while True:
            random_number = random.randint(1000, 9999)
            if random_number not in self.sequence_number:
                self.sequence_number.add(random_number)
                return random_number

    @staticmethod
    def text(paragraphs, old_text, new_text):
        if old_text in paragraphs.text:
            paragraphs.text = paragraphs.text.replace(old_text, new_text)

    def invoice(self):
        self.current_invoice_number = self.number()
        doc = docx.Document("invoice.docx")
        try:
            replacements = {
                "[name]": self.name_entry.get().upper(),
                "[number]": self.number_entry.get(),
                "[email]": self.email_entry.get(),
                "[date]": datetime.datetime.now().strftime("%d/%m/%Y"),
                "[time]": datetime.datetime.now().strftime("%H:%M:%S"),
                "[inv]": self.sequence_id,
                "[year]": datetime.datetime.now().strftime("%Y"),
                "[num]": str(self.current_invoice_number)
            }
        except KeyError:
            messagebox.showinfo("Error", "Please fill all the required fields")
            return

        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.text(paragraph, old_text, new_text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.text(paragraph, old_text, new_text)

        saved_path = filedialog.asksaveasfilename(defaultextension='.pdf', filetypes=[('PDF document', '*.pdf')])
        if not saved_path:
            messagebox.showinfo("Cancelled", "Invoice creation cancelled.")
            return

        temp_path = "filled_temp.docx"
        doc.save(temp_path)
        try:
            convert(temp_path, os.path.dirname(saved_path))
            final_pdf_path = os.path.join(os.path.dirname(saved_path), 'filled_temp.pdf')
            os.rename(final_pdf_path, saved_path)
            os.remove(temp_path)
            messagebox.showinfo("Success", f"Invoice created successfully at:\n{saved_path}")
        except Exception as e:
            messagebox.showinfo("Error", f"An error occurred: {e}")
            os.remove(temp_path)
        self.save_to_excel()

    def save_to_excel(self):
        name = self.name_entry.get().strip()
        number = self.number_entry.get().strip()
        email = self.email_entry.get().strip()
        if not name or not number or not email:
            messagebox.showerror("Validation Error", "All fields are required!")
            return

        file_path = "../excel/USER_DATA.xlsx"
        try:
            if os.path.exists(file_path):
                workbook = load_workbook(file_path)
                sheet = workbook.active
            else:
                workbook = Workbook()
                sheet = workbook.active
                sheet.append(["Name", "Number", "Email", "Date", "Time", "INVOICE NUMBER"])

            sheet.append([name, number, email,
                          datetime.datetime.now().strftime("%d/%m/%Y"),
                          datetime.datetime.now().strftime("%H:%M:%S"),
                          f"{self.sequence_id}-{datetime.datetime.now().strftime('%Y')}-{str(self.current_invoice_number)}"])

            workbook.save(file_path)
            workbook.close()
            messagebox.showinfo("Success", f"User data saved successfully to {file_path}!")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while saving to Excel: {e}")


if __name__ == "__main__":
    app = Application()
